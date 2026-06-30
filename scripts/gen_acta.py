#!/usr/bin/env python3
"""Genera el Acta de Entrega del sitio web The Tantric Garden en formato .docx."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x6B, 0x4E, 0x2E)   # marrón cálido de la marca
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# --- estilos base ---
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

def hl(run):
    """Resalta un run en amarillo (campo POR COMPLETAR)."""
    rPr = run._element.get_or_add_rPr()
    h = OxmlElement('w:highlight')
    h.set(qn('w:val'), 'yellow')
    rPr.append(h)
    return run

def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = ACCENT
    p.space_before = Pt(14)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4'); bottom.set(qn('w:color'), '6B4E2E')
    pbdr.append(bottom); pPr.append(pbdr)
    return p

def field(label, value, todo=False, secure=False):
    p = doc.add_paragraph()
    rl = p.add_run(f'{label}: ')
    rl.bold = True
    if secure:
        rs = p.add_run('🔒 Se entrega por canal seguro (no se incluye en este documento)')
        rs.font.color.rgb = GREY
        rs.italic = True
    elif todo:
        hl(p.add_run(value if value else 'POR COMPLETAR'))
    else:
        p.add_run(value)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)
    return p

def bullet(text, check=False):
    p = doc.add_paragraph(style='List Bullet')
    if check:
        rc = p.add_run('✔  '); rc.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32); rc.bold = True
    p.add_run(text)
    return p

# ============ ENCABEZADO ============
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt = title.add_run('ACTA DE ENTREGA DEL SITIO WEB')
rt.bold = True; rt.font.size = Pt(20); rt.font.color.rgb = ACCENT

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run('The Tantric Garden  ·  tantricmallorca.com')
rs.font.size = Pt(12); rs.font.color.rgb = GREY

# ============ 1. INFORMACIÓN GENERAL ============
h1('1. Información general')
field('Nombre del proyecto', 'Sitio web corporativo — The Tantric Garden')
field('Cliente', 'The Tantric Garden — Palma de Mallorca (España)')
field('Responsable de la entrega', 'David ', todo=True)
field('Fecha de entrega', '16 de junio de 2026')
field('URL del sitio web', 'https://tantricmallorca.com')

# ============ 2. ALCANCE ENTREGADO ============
h1('2. Alcance entregado')
bullet('Diseño y desarrollo del sitio web a medida (sin plantilla y sin CMS), según lo aprobado.', check=True)
bullet('Implementación y publicación del sitio en producción (hosting Vercel).', check=True)
bullet('Configuración responsive (adaptado a móvil, tablet y escritorio).', check=True)
bullet('Integración de analítica: Google Analytics 4 con Consent Mode v2 y banner de cookies (RGPD).', check=True)
bullet('Botón / llamada a la acción de contacto directo por WhatsApp.', check=True)
bullet('SEO técnico: metaetiquetas, sitemap, favicons y renderizado previo (SSR/prerender) de las páginas.', check=True)
bullet('Pruebas funcionales realizadas en producción.', check=True)

# ============ 3. CREDENCIALES ENTREGADAS ============
h1('3. Credenciales y accesos entregados')
note = doc.add_paragraph()
rn = note.add_run('Nota de seguridad: por buenas prácticas, las contraseñas NO se incluyen en este documento. '
                  'Se entregan al cliente por un canal seguro (gestor de contraseñas o mensaje cifrado aparte).')
rn.italic = True; rn.font.size = Pt(9.5); rn.font.color.rgb = GREY

doc.add_paragraph().add_run('Hosting').bold = True
field('Proveedor', 'Vercel (vercel.com)')
field('Usuario / cuenta', 'Vinculada a la cuenta de GitHub del proyecto', todo=True)
field('Correo asociado', 'tantricgardenmallorca@gmail.com')
field('Contraseña', '', secure=True)

doc.add_paragraph().add_run('Dominio').bold = True
field('Proveedor', 'DonDominio (dondominio.com)')
field('Usuario o titular', 'Titular de la cuenta DonDominio', todo=True)
field('Correo asociado', 'tantricgardenmallorca@gmail.com')
field('Contraseña', '', secure=True)

doc.add_paragraph().add_run('Repositorio de código fuente').bold = True
field('Plataforma', 'GitHub')
field('URL del repositorio', 'https://github.com/tantricgardenmallorca/tantricmallorca')
field('Usuario / organización', 'tantricgardenmallorca')

doc.add_paragraph().add_run('CMS o administrador del sitio').bold = True
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.3)
p.add_run('No aplica. ').bold = True
p.add_run('El sitio fue desarrollado a medida (Vite + React); no existe un panel tipo WordPress/Squarespace. '
          'Los cambios de contenido se realizan en el código del repositorio y se publican automáticamente '
          'en Vercel al subirlos a GitHub.')

doc.add_paragraph().add_run('Herramientas adicionales').bold = True
field('Google Analytics 4', 'Measurement ID G-PC7VN218GR — propiedad "The Tantric Garden — Web"')
field('Google Tag Manager', 'No se utiliza (la medición se hace con gtag.js directo)')
field('Search Console', 'Propiedad verificada: sc-domain:tantricmallorca.com')
field('Correo corporativo', 'tantricgardenmallorca@gmail.com')
field('Perfil de Empresa en Google', 'Ficha "The Tantric Garden" (Google Business Profile), administrada con el mismo correo')

# ============ 4. RENOVACIONES ============
h1('4. Renovaciones')

table = doc.add_table(rows=1, cols=2)
table.style = 'Light List Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.LEFT
hdr = table.rows[0].cells
hdr[0].paragraphs[0].add_run('Concepto').bold = True
hdr[1].paragraphs[0].add_run('Detalle').bold = True

rows = [
    ('Dominio — tantricmallorca.com', 'Registrado en DonDominio. Renovación ANUAL.'),
    ('Fecha de vencimiento del dominio', '__POR COMPLETAR__ (consultar en DonDominio → Mis dominios)'),
    ('Hosting — Vercel', 'Plan __POR COMPLETAR__ (Hobby/gratuito o Pro). En plan Hobby no hay renovación ni costo anual; el sitio sigue publicado mientras la cuenta esté activa.'),
    ('Responsable de las renovaciones', '__POR COMPLETAR__'),
    ('Método de pago asociado (opcional)', '__POR COMPLETAR__'),
]
for concepto, detalle in rows:
    cells = table.add_row().cells
    cells[0].paragraphs[0].add_run(concepto).bold = True
    if '__POR COMPLETAR__' in detalle:
        before, rest = detalle.split('__POR COMPLETAR__', 1)
        para = cells[1].paragraphs[0]
        if before: para.add_run(before)
        hl(para.add_run('POR COMPLETAR'))
        if rest: para.add_run(rest)
    else:
        cells[1].paragraphs[0].add_run(detalle)

doc.add_paragraph()
exp = doc.add_paragraph()
exp.add_run('Cómo renovar el dominio (DonDominio): ').bold = True
steps = [
    'Entrar a https://www.dondominio.com con la cuenta titular (correo y contraseña entregados por canal seguro).',
    'Ir a "Mis dominios" y seleccionar tantricmallorca.com.',
    'Pulsar "Renovar" y elegir el periodo (normalmente 1 año).',
    'Confirmar el pago con el método asociado. Se recomienda renovar antes de la fecha de vencimiento para no perder el dominio.',
    'Opcional: activar la "renovación automática" para que se cobre solo cada año.',
]
for i, s in enumerate(steps, 1):
    pp = doc.add_paragraph(style='List Number')
    pp.add_run(s)

note2 = doc.add_paragraph()
rn2 = note2.add_run('El hosting en Vercel (si está en plan Hobby) no requiere renovación de pago; basta con mantener la '
                    'cuenta de GitHub y Vercel activas. Si en algún momento se cambia a plan Pro, la facturación sería mensual.')
rn2.italic = True; rn2.font.size = Pt(9.5); rn2.font.color.rgb = GREY

# ============ 5. ARCHIVOS Y ACTIVOS ============
h1('5. Archivos y activos entregados')
bullet('Logo e identidad visual (versiones clara y oscura en SVG, favicons en todos los tamaños).', check=True)
bullet('Fotografías del local y del equipo utilizadas en el sitio (incluidas en el repositorio).', check=True)
bullet('Código fuente completo del sitio (repositorio de GitHub indicado arriba).', check=True)
bullet('Este documento de entrega, con instructivo de renovación del dominio.', check=True)
bullet('Copia de seguridad del sitio: el propio repositorio de GitHub actúa como respaldo versionado.', check=True)

# ============ 6. GARANTÍA / SOPORTE ============
h1('6. Garantía o soporte posterior')
p = doc.add_paragraph()
p.add_run('Se incluye soporte por ')
hl(p.add_run('30'))
p.add_run(' días a partir de la fecha de entrega para la corrección de errores relacionados con el alcance '
          'inicialmente aprobado. Nuevos requerimientos o desarrollos adicionales serán cotizados por separado.')

# ============ 7. OBSERVACIONES ============
h1('7. Observaciones')
doc.add_paragraph('El cliente declara haber revisado el sitio web y recibido los accesos e información relacionados '
                  'con el proyecto, manifestando su conformidad con la entrega realizada.')

# ============ 8. APROBACIÓN ============
h1('8. Aprobación')
sign = doc.add_table(rows=1, cols=2)
sign.alignment = WD_TABLE_ALIGNMENT.CENTER
c = sign.rows[0].cells
for cell, rol in zip(c, ['CLIENTE', 'RESPONSABLE DE LA ENTREGA']):
    cell.paragraphs[0].add_run(rol).bold = True
    cell.add_paragraph('\n\n______________________________')
    cell.add_paragraph('Nombre:')
    cell.add_paragraph('Firma:')
    cell.add_paragraph('Fecha:')

doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = foot.add_run('Los campos resaltados en amarillo deben completarse antes de la firma.')
rf.italic = True; rf.font.size = Pt(9); rf.font.color.rgb = GREY

out = '/Users/kevin/Desktop/Acta-de-Entrega-Tantric-Garden.docx'
doc.save(out)
print('Guardado:', out)
